"""
api/routers/documents.py
-------------------------
Prefix: /admin
Routes:
  POST   /admin/upload-doc      — parse & store PDF/DOCX/TXT/CSV/XLSX
  POST   /admin/add-url         — scrape a URL and store as knowledge doc
  POST   /admin/refresh-urls    — re-scrape all URL-type docs (called by scheduler)
  GET    /admin/docs            — list active documents
  DELETE /admin/doc/{doc_id}    — soft-delete a document
"""
import io
import logging
from typing import List

import PyPDF2
import pandas as pd
import httpx
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from fastapi import APIRouter, Depends, File, HTTPException, Query, UploadFile
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from ..deps import get_tenant_db
from ....database import KnowledgeDocument, get_tenant_limits
from ...schemas.models import DocumentResponse
from ...utils.formatters import format_utc

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/admin", tags=["Documents"])

MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024
ALLOWED_EXTENSIONS = {".pdf", ".txt", ".csv", ".xlsx", ".xls", ".docx"}


class AddUrlRequest(BaseModel):
    url: str


def _scrape_url(url: str) -> str:
    try:
        with httpx.Client(timeout=15, follow_redirects=True) as client:
            response = client.get(url, headers={"User-Agent": "Mozilla/5.0 AI-Training-Bot/1.0"})
            response.raise_for_status()
    except httpx.HTTPStatusError as e:
        raise HTTPException(status_code=400, detail=f"URL returned HTTP {e.response.status_code}: {url}")
    except httpx.RequestError as e:
        raise HTTPException(status_code=400, detail=f"Could not reach URL: {str(e)}")

    soup = BeautifulSoup(response.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside", "form"]):
        tag.decompose()
    text = soup.get_text(separator="\n", strip=True)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return "\n".join(lines)


def _parse_file(filename: str, content_bytes: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        reader = PyPDF2.PdfReader(io.BytesIO(content_bytes))
        text = ""
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
        return text
    elif lower.endswith(".docx"):
        doc = DocxDocument(io.BytesIO(content_bytes))
        return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    elif lower.endswith(".txt"):
        return content_bytes.decode("utf-8", errors="replace")
    elif lower.endswith(".csv"):
        return content_bytes.decode("utf-8", errors="replace")
    elif lower.endswith((".xlsx", ".xls")):
        df = pd.read_excel(io.BytesIO(content_bytes))
        return df.to_string()
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file format: {filename}")


@router.post("/upload-doc", response_model=DocumentResponse)
async def upload_document(
    tenant_id: str = Query(...),
    file: UploadFile = File(...),
    db: Session = Depends(get_tenant_db),
):
    ext = "." + file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
        )

    limits = get_tenant_limits(tenant_id)
    doc_count = db.query(KnowledgeDocument).filter(
        KnowledgeDocument.tenant_id == tenant_id, KnowledgeDocument.is_active == True
    ).count()
    if doc_count >= limits.get("docs", 5):
        raise HTTPException(
            status_code=403,
            detail=f"Document storage limit reached ({limits.get('docs')}). Please upgrade.",
        )

    content_bytes = await file.read()
    if len(content_bytes) > MAX_FILE_SIZE_BYTES:
        raise HTTPException(
            status_code=413,
            detail=f"File too large ({len(content_bytes) // 1024} KB). Maximum allowed size is 10 MB.",
        )

    try:
        content = _parse_file(file.filename, content_bytes)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse file: {str(e)}")

    if not content.strip():
        raise HTTPException(status_code=400, detail="File appears to be empty or could not be parsed.")

    doc = KnowledgeDocument(
        tenant_id=tenant_id,
        filename=file.filename,
        content=content,
        file_type=ext.lstrip("."),
        file_size_bytes=len(content_bytes)
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return {
        "id": doc.id, "filename": doc.filename, "file_type": doc.file_type,
        "is_active": doc.is_active, "created_at": format_utc(doc.created_at),
        "file_size_bytes": doc.file_size_bytes
    }


@router.post("/add-url", response_model=DocumentResponse)
async def add_url_document(
    tenant_id: str = Query(...),
    payload: AddUrlRequest = ...,
    db: Session = Depends(get_tenant_db),
):
    url = payload.url.strip()
    if not url.startswith(("http://", "https://")):
        raise HTTPException(status_code=400, detail="URL must start with http:// or https://")

    limits = get_tenant_limits(tenant_id)
    doc_count = db.query(KnowledgeDocument).filter(
        KnowledgeDocument.tenant_id == tenant_id, KnowledgeDocument.is_active == True
    ).count()
    if doc_count >= limits.get("docs", 5):
        raise HTTPException(
            status_code=403,
            detail=f"Document storage limit reached ({limits.get('docs')}). Please upgrade.",
        )

    existing = db.query(KnowledgeDocument).filter(
        KnowledgeDocument.tenant_id == tenant_id,
        KnowledgeDocument.filename == url,
        KnowledgeDocument.is_active == True,
    ).first()

    content = _scrape_url(url)

    if not content.strip():
        raise HTTPException(status_code=400, detail="Could not extract any text content from that URL.")

    if existing:
        existing.content = content
        existing.file_size_bytes = len(content.encode('utf-8'))
        db.commit()
        db.refresh(existing)
        return {
            "id": existing.id, "filename": existing.filename, "file_type": existing.file_type,
            "is_active": existing.is_active, "created_at": format_utc(existing.created_at),
            "file_size_bytes": existing.file_size_bytes
        }

    doc = KnowledgeDocument(
        tenant_id=tenant_id,
        filename=url,
        content=content,
        file_type="url",
        file_size_bytes=len(content.encode('utf-8'))
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return {
        "id": doc.id, "filename": doc.filename, "file_type": doc.file_type,
        "is_active": doc.is_active, "created_at": format_utc(doc.created_at),
        "file_size_bytes": doc.file_size_bytes
    }


@router.post("/refresh-urls")
async def refresh_url_documents(tenant_id: str = Query(...), db: Session = Depends(get_tenant_db)):
    url_docs = db.query(KnowledgeDocument).filter(
        KnowledgeDocument.tenant_id == tenant_id,
        KnowledgeDocument.file_type == "url",
        KnowledgeDocument.is_active == True,
    ).all()

    refreshed, failed = 0, 0
    for doc in url_docs:
        try:
            new_content = _scrape_url(doc.filename)
            if new_content.strip():
                doc.content = new_content
                doc.file_size_bytes = len(new_content.encode('utf-8'))
                refreshed += 1
        except Exception as e:
            logger.warning(f"[URL Refresh] Failed for {doc.filename}: {e}")
            failed += 1

    db.commit()
    return {"refreshed": refreshed, "failed": failed}


@router.get("/docs", response_model=List[DocumentResponse])
async def get_documents(tenant_id: str = Query(...), db: Session = Depends(get_tenant_db)):
    docs = db.query(KnowledgeDocument).filter(
        KnowledgeDocument.tenant_id == tenant_id, KnowledgeDocument.is_active == True
    ).all()
    return [
        {
            "id": d.id, "filename": d.filename, "file_type": d.file_type,
            "is_active": d.is_active, "created_at": format_utc(d.created_at),
            "file_size_bytes": d.file_size_bytes
        }
        for d in docs
    ]


@router.delete("/doc/{doc_id}")
async def delete_document(doc_id: str, tenant_id: str = Query(...), db: Session = Depends(get_tenant_db)):
    doc = db.query(KnowledgeDocument).filter(
        KnowledgeDocument.id == doc_id, KnowledgeDocument.tenant_id == tenant_id
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    doc.is_active = False
    db.commit()
    return {"message": "Document deleted"}
